"""
Import pèlerins from Word (.docx) files into the database.
Parses vol info (2 aller legs, 2 retour legs) from paragraphs and pilgrim data from tables.
"""
import sys
import os
import re
from docx import Document
from sqlmodel import Session
from database import engine, init_db
from models import Pelerin


def extract_vol_info(doc):
    """Extract vol number and flight legs from paragraphs."""
    vol_num = ""
    aller_legs = []
    retour_legs = []
    current_section = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or text == 'gggggggggggggggggggg':
            continue

        upper = text.upper()

        # Vol number
        vol_match = re.search(r'VOL\s*(\d+)', upper)
        if vol_match and not vol_num:
            vol_num = f"VOL {vol_match.group(1)}"

        # Detect section
        if 'ALLEZ' in upper or 'ALLER' in upper:
            current_section = 'aller'
        elif 'RETOUR' in upper:
            current_section = 'retour'

        # Extract flight leg: "ET 934 18 MAI ABJ – ADD 12H00MN – 21H00MN"
        leg_match = re.search(r'(ET\s*\w+)\s+(\d{1,2}\s+\w+)\s+(.+)', text, re.IGNORECASE)
        if leg_match:
            flight_num = leg_match.group(1).strip()
            date = leg_match.group(2).strip()
            route = leg_match.group(3).strip()
            leg_str = f"{flight_num} — {date} — {route}"

            if current_section == 'aller':
                aller_legs.append(leg_str)
            elif current_section == 'retour':
                retour_legs.append(leg_str)

    return vol_num, aller_legs, retour_legs


def identify_column(header_text: str) -> str | None:
    h = header_text.lower().strip()
    h = h.replace("é", "e").replace("è", "e").replace("ê", "e")
    h = re.sub(r'[^a-z\s]', ' ', h).strip()
    h = re.sub(r'\s+', ' ', h)

    if h == "nom":
        return "nom"
    if h == "prenom" or h == "prenoms":
        return "prenom"
    if "passport" in h or "passeport" in h:
        return "numero_passeport"
    if h == "statut":
        return "statut"
    return None


def find_column_mapping(headers):
    mapping = {}
    for i, h in enumerate(headers):
        field = identify_column(h)
        if field:
            mapping[i] = field
    return mapping


def import_docx(filepath, session):
    doc = Document(filepath)
    filename = os.path.basename(filepath)

    vol_num, aller_legs, retour_legs = extract_vol_info(doc)
    
    vol_aller_1 = aller_legs[0] if len(aller_legs) > 0 else ""
    vol_aller_2 = aller_legs[1] if len(aller_legs) > 1 else ""
    vol_retour_1 = retour_legs[0] if len(retour_legs) > 0 else ""
    vol_retour_2 = retour_legs[1] if len(retour_legs) > 1 else ""

    print(f"  ✈️  {vol_num}")
    print(f"     Aller 1: {vol_aller_1}")
    print(f"     Aller 2: {vol_aller_2}")
    print(f"     Retour 1: {vol_retour_1}")
    print(f"     Retour 2: {vol_retour_2}")

    total = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_map = find_column_mapping(headers)

        if not col_map:
            continue

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue

            data = {}
            for idx, field in col_map.items():
                if idx < len(cells):
                    data[field] = cells[idx]

            nom = data.get("nom", "").strip()
            if not nom:
                continue

            pelerin = Pelerin(
                nom=nom,
                prenom=data.get("prenom", "").strip(),
                numero_passeport=data.get("numero_passeport", "").strip(),
                numero_vol=vol_num,
                vol_aller_1=vol_aller_1,
                vol_aller_2=vol_aller_2,
                vol_retour_1=vol_retour_1,
                vol_retour_2=vol_retour_2,
                statut=data.get("statut", "").strip(),
                source_fichier=filename,
            )
            session.add(pelerin)
            total += 1

    session.commit()
    return total


def main():
    if len(sys.argv) < 2:
        print("Usage: python import_word.py <file1.docx> [file2.docx ...]")
        sys.exit(1)

    init_db()

    with Session(engine) as session:
        for filepath in sys.argv[1:]:
            print(f"\n📄 Import de: {filepath}")
            if not os.path.exists(filepath):
                print(f"  ❌ Fichier introuvable")
                continue
            count = import_docx(filepath, session)
            print(f"  ✅ {count} pèlerins importés")


if __name__ == "__main__":
    main()
